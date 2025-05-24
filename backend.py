# backend.py

import os
import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image
import tempfile

from langchain.embeddings import OllamaEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

CHROMA_PATH = "./chroma_db"

embedding = OllamaEmbeddings(model="phi3:mini")

def extract_text(file):
    extension = os.path.splitext(file.name)[1].lower()

    if extension == ".pdf":
        text = ""
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text
                else:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    text += ocr_text
        return text

    elif extension == ".docx":
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])

    elif extension == ".txt":
        return file.read().decode("utf-8")

    else:
        raise ValueError("Unsupported file type")

def extract_text_from_pdf(path):
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text += page_text
            else:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                text += ocr_text
    return text

def extract_text_from_docx(path):
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_preview_from_file(filepath, filename=None):
    if filepath.endswith(".pdf"):
        text = extract_text_from_pdf(filepath)
    elif filepath.endswith(".docx"):
        text = extract_text_from_docx(filepath)
    elif filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError("Unsupported file format")
    
    return text[:1000] if text else "(No preview available)"

def create_vectorstore_from_uploaded_file(filepath, filename=None):
    text = ""

    if filepath.endswith(".pdf"):
        text = extract_text_from_pdf(filepath)
    elif filepath.endswith(".docx"):
        text = extract_text_from_docx(filepath)
    elif filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError("Unsupported file format")

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    documents = splitter.create_documents([text])
    vectorstore = Chroma.from_documents(documents, embedding, persist_directory=CHROMA_PATH)
    vectorstore.persist()

def query_rag(user_query):
    vectorstore = Chroma(persist_directory=CHROMA_PATH, embedding_function=embedding)
    retriever = vectorstore.as_retriever()
    docs = retriever.get_relevant_documents(user_query)

    context = "\n".join([doc.page_content for doc in docs])

    prompt = f"""
You are a helpful and knowledgeable assistant. Use the following context extracted from a document to accurately and clearly answer the user's question.

Only use information from the context. If you are unsure or the context doesn’t provide enough information, say so clearly.

Context:
{context}

Question:
{user_query}

Answer:"""

    print("Using model: phi3:mini")

    import requests
    response = requests.post("http://localhost:11434/api/generate", json={
        "model": "phi3:mini",
        "prompt": prompt,
        "stream": False
    })

    return response.json().get("response", "No response from model.")
